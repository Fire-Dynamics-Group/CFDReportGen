import datetime
import sys
import os

print("path: ",os.path.dirname(sys.executable))
import PySimpleGUI as sg
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches

# https://docxtpl.readthedocs.io/en/latest/


from helper_functions import round_to, find_all_files_of_type
from scenarios_object import create_scenario_object
from scen_object_helper_functions import return_fds_version
from report_gen_helper_functions import scen_results_values
from hrr_graph import run_CFD_charts
from validate import validate_form, generate_error_message, scenario_types
from final_report import create_report

document_name = "Template CFD Report.docx"
# check if file exists, if not use binary
document_path = f"{document_name}" # Path(__file__).parent /"CFD Word Template"/document_name
if os.path.exists(document_path):
    doc = DocxTemplate(document_path)
# else:
#         # Decode the base64 string to bytes
#     TEMPLATE_BASE64 = ''
#     template_bytes = base64.b64decode(TEMPLATE_BASE64)

#     # Create a BytesIO object and load it into DocxTemplate
#     buffer = io.BytesIO(template_bytes)
#     tpl = DocxTemplate(buffer)
# #     # Decode the base64 string back to bytes
# #     def decode_base64_to_bytes(base64_string):
# #         return base64.b64decode(base64_string.encode('utf-8'))

# #     # Create a Word Document from the embedded base64 string
# #     def create_word_from_template():
# #         template_bytes = decode_base64_to_bytes(TEMPLATE_BASE64)
# #         buffer = io.BytesIO(template_bytes)
# #         document = DocxTemplate(buffer)

# #         # buffer.close()
# #         # Do something with the document, e.g., save or add content
# #         return document
# #     doc = create_word_from_template()
#     context = {
#     'my_placeholder': 'This is replaced text',
#     # add more placeholders here
# }

# # # Render the template
#     tpl.render(context)
#     doc.render(context)





''' 
    TODO: have gui popup first
    allow path to be entered by user
'''
# Delete_row_in_table(2, -1)

# print(document.tables[0])

# for tables in document.tables:
#     for rows in tables.rows:
#         for cells in rows.cells:
#             print(cells.text)

# document = Document(document_path)
# get all styles from doc
# may need to get index of paragraph and match to styles each time
# search for curly braces??
# styles = document.styles
# paragraph_styles = [
#     s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
# ]





def create_inline_image(image_file, template=doc):
    return InlineImage(template, image_descriptor=image_file, width=Inches(6), height=Inches(4)) # can add relative path entry box??
# charts= {}
# for chart in chart_names:
#     charts[chart] = create_inline_image(f"png_charts/{chart}", template=doc)


today = datetime.datetime.today()

'''TODO: move gui to own script -> allowing path to be entered manually '''
# GUI input boxes etc in layout
layout = [
       [sg.Text("Path to runs' root directory:"), sg.Input(key="PATH", do_not_clear=True)],        
       [sg.Text("Client name:"), sg.Input(key="CLIENT_NAME", do_not_clear=True)],
        [sg.Text("Project name:"), sg.Input(key="PROJECT_NAME", do_not_clear=True)],
        [sg.Text("Project Locations:"), sg.Input(key="PROJECT_LOCATION", do_not_clear=True, size=(20,1))],
        [sg.Text("Senior's email prefix:"), sg.Input(key="EMAIL_PREFIX", do_not_clear=True, size=(20,1))],
        [sg.Text("Extended Travel Distances:"), 
        sg.Radio('True', group_id="EXTENDED_TRAVEL", key="HAS_EXTENDED_TRAVEL",default=True),
        sg.Radio('False', group_id="EXTENDED_TRAVEL", key="NO_EXTENDED_TRAVEL")],
        # if values["HAS_EXTENDED_TRAVEL"]:
        [sg.Text("Max Travel Distance:"), sg.Input(key="MAX_TD", do_not_clear=True, size=(20,1)), sg.Text("m")],
        # guidance doc
        [sg.Text("Guidance Doc:"), 
        sg.Radio('BS9991', group_id="GUIDANCE", key="BS9991",default=True),
        sg.Radio('ADB', group_id="GUIDANCE", key="ADB")],
        # [sg.Text("Extended Travel Distances:"), sg.Listbox((["True", "False"]), size=(20,4), enable_events=False, key='HAS_EXTENDED_TRAVEL')],
        [sg.Button("Create Report"), sg.Exit()], 
]

window = sg.Window("Report Generator", layout, element_justification="right")

while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED or event == "Exit":
        break
    if event == "Create Report":
        ''' 
        TODO: verify input and template file etc 
        check path
        '''
        if values['PATH']:
            values['PATH'] = r"{}".format(values['PATH'])
        is_valid, values_invalid = validate_form(values)
        files_error_message = ""
        if is_valid == False:
            error_message = generate_error_message(values_invalid)
            sg.popup_error(error_message,title="Form Input Error")
        else:        
            values["TODAYS_DATE"] = today.strftime("%d-%m-%Y")
            # assumed if first scenario has sprinklers all do 
            # chart_names = [ f for f in listdir("png_charts") if isfile(join("png_charts", f)) ]

            path_to_directory="graph_generation"
            # TODO: get path from user - perhaps initial input before further one
            # C:\Users\IanShaw\Dropbox\Projects CFD\25. Claridges\Runs
            # path_to_root_directory = Path(r"C:\Users\IanShaw\Dropbox\Projects CFD\26. Breams Building\Runs")
            path_to_root_directory = f"{values['PATH']}"
            # path_to_root_directory = Path(r"C:\Users\IanShaw\Dropbox\Projects CFD\9. 100 Avenue Road\Jan 2023 Corridor Models")
            # path_to_root_directory = Path(r"C:\Users\IanShaw\Dropbox\Projects CFD\22. Sweet Street\Resi\Final")
            # path_to_root_directory = Path(r"C:\Users\IanShaw\Dropbox\Projects CFD\1. Graph Generation\Test Cases\Test4")
            
            '''
            # TODO: check folder structure
            # if no subfolder -> message to user to add subfolder
            # if fds file not named appropriately -> message to user to rename
            # 
            '''
            # have path_to_dir changed if required
            # trial_path = f'{path_to_root_directory}/{scenario_name}'

            scenarios_object, scenario_names, FSA_scenarios, MoE_scenarios, error_list = create_scenario_object(path_to_directory=path_to_root_directory)
            if len(error_list) > 0:
                sg.popup_error("Error", '\n\n'.join(error_list))
            # have popup -> x MOE runs, y FSA runs -> ask to continue or rename folders with 'FSA'
            files_error_message = scenario_types(FSA_scenarios, MoE_scenarios)
            # sg.popup_error(files_error_message,title="Form Input Error")
            sg.popup("Scenarios Found:", files_error_message)
        if is_valid and len(error_list) == 0:        
            from render_report import render_logic
            render_logic(values, today, path_to_root_directory, scenarios_object, scenario_names, MoE_scenarios, FSA_scenarios)
            # jinja should output x amount of scenario loops?
            # TODO: use output path above for images to be saved to
            if os.path.isdir("outputReports"):
                new_dir_path = f"outputReports/{values['PROJECT_NAME']}" #Path(__file__).parent / "outputReports"/f"{values['PROJECT_NAME']}"
            else:
                new_dir_path = f"{values['PROJECT_NAME']}"
            
            os.mkdir(new_dir_path)

            run_CFD_charts(path_to_root_directory, scenario_names, new_dir_path)
            # TODO: output charts into report
            # scope images in folder
            def insert_charts(new_dir_path):
                # look at charts in folder
                chart_names = find_all_files_of_type(new_dir_path, suffix=".png")
                prefixes = [f.split("_")[0] for f in chart_names]
                unique_prefixes = list(set(prefixes))
                # either have a million figure tags already in 
                # loop through scenarios - need unique prefixes before underscore from filename
                # trial insert first chart
                charts= {}
                for chart in chart_names:
                    charts[chart] = create_inline_image(f"{new_dir_path}/{chart}", template=doc)
                # filter for each scenario
                # input in order: hrr, vis, (vs time moe/vs distance fsa)temp, zslice, max p drop
                # then 
                i = 0
                for i in range(len(unique_prefixes)):
                    current_scen_chart_names = [f for f in chart_names if unique_prefixes[i] in f]
                    def set_chart(chart_type = "hrr"):
                        current_chart_list = [f for f in current_scen_chart_names if chart_type in f.lower()]
                        if len(current_chart_list) > 0:
                            current_chart = current_chart_list[0]
                        # current_chart = current_list
                            values[f"SCEN_{i+1}_{chart_type.upper()}_CHART"] = charts[current_chart]
                    set_chart("hrr")
                    set_chart("vis")
                    set_chart("temp")
                    set_chart("pres")
                print("test")

                
            # # should have scenario name/number in title of charts
            # values[f"SCENARIO_{scenario_index}_PRESSURE_CC_MOE"] = charts["Pressure cc_pres__chart.png"] 
            # values["SCENARIO_1_HRR_CC_MOE"] = charts["hrr_chart.png"]
            # values["SCENARIO_1_VELOCITY_CC_MOE"] = charts["Velocity cc_vel__chart.png"]
            # values["SCENARIO_1_VISIBILITY_CC_MOE"] = charts["Visibility cc_vis__chart.png"]
            insert_charts(new_dir_path)


            # Render the template, save new word document & inform user
            doc.render(values)



            output_path = f"{new_dir_path}/{values['PROJECT_NAME']}-CFD Report.docx"
            doc.save(output_path)


            create_report(output_path, ref_order, fds_version, scenario_names, scenarios_object, MoE_scenarios, FSA_scenarios)
            sg.popup("File saved", f"File has been saved here: {output_path}")
            # then apply docx to tables etc
            # TODO: docx separate script, output_path sent as parameter, import object
 
            startup_path = rf'{os.path.abspath(os.getcwd())}/{output_path}'
            os.startfile(startup_path)

            window.close()


    # Below for without gui

    # {{PROJECT NAME}} {{TODAYS DATE}} {{CLIENT NAME}}

    # project_name = "Test Resi Project"
    # client_name = "Test Client"

    # context = {
    #     "PROJECT_NAME": project_name,
    #     "TODAYS_DATE": today.strftime("%d-%m-%Y"),
    #     "CLIENT_NAME": client_name
    # }
    # doc.render(context)
    # doc.save(Path(__file__).parent / f"{project_name}-Fire Dynamics.docx")