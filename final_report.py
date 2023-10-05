from docx import Document
from constants import font_name_light
from docx.shared import Pt, RGBColor
from pathlib import Path

from report_gen_helper_functions import scen_results_values
# TODO: move to misc utils
    # TODO: pull reference order from jinja tags in document -> then remove adb/bs9991




def Delete_row_in_table(table, row, document):
    if type(table) == int:
        table_object = document.tables[table]
    else:
        table_object = table
    table_object._tbl.remove(table_object.rows[row]._tr)

# below not needed to test charts etc
def create_report(output_path, fds_version, ref_order, scenario_names, scenarios_object, MoE_scenarios, FSA_scenarios):
    document = Document(output_path)
            # loop below through each row and column
    def replace_table_cell_content(cell, replacement_text, is_bold=False, alignment=1):
        # later make object with fonts etc
        cell.text = replacement_text
        paragraphs = cell.paragraphs
        paragraphs[0].alignment = alignment # 1 = centered
        # for run in paragraphs.runs:
        run = paragraphs[0].runs
        font = run[0].font
        font.size = Pt(9) # pull from object
        font.name = font_name_light
        font.color.rgb = RGBColor(64,64,64) #gray color # to be pulled from object
        font.bold = is_bold

    # TODO: references: number superscript in report using Jinjs and 
    # TODO: have in order in references section 
    # TODO: pull references from csv
    # below would change for other users!!!
    ref_repo_file = Path(r'C:\Users\IanShaw\Fire Dynamics Group Limited\Research - Ian\CFD Report Generator\other\references.csv')
    # TODO: remove unused lines
    with open(ref_repo_file, "r+", encoding="utf8") as f:
        ref_repo_list = f.readlines()[1:]
    split_repo_list = [f.split(",") for f in ref_repo_list] # id, title, ...ref_info
    document_text = [para.text for para in document.paragraphs]
    ref_table_paras = [para for para in document.paragraphs if "REF_" in para.text]

    # reference still accessible using below
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    
    # TODO: follow order of ref_order list
    # TODO: remove unused rows
    for ref_table_index in range(len(ref_order)):
    # ref_table_index = 0
        # TODO: find id from csv
        target_id = ref_order[ref_table_index]
        # loop through lines -> find same id 
        for csv_index in range(len(split_repo_list)):
            test_line = split_repo_list[csv_index]
            id, ref_title, *ref_info = test_line
            # TODO: if fds -> insert version used from fds_version
            if target_id == id:
                # target_line = test_line
                ref_info = ','.join(str(item) for item in ref_info)
                ref_info.strip('\n')
                para = ref_table_paras[ref_table_index]
                para.alignment = 0
                para.clear()
                if id == "FDS":
                    ref_title = f'FDS Version {fds_version}'
                run = para.add_run(ref_title.replace('"', ''))
                run.bold = True
                para.add_run(" ")
                para.add_run(ref_info.replace('"', ''))
                break
    # for index in range(10):
    #     delete_paragraph(ref_table_paras[-(index+1)])
    # TODO: remove excess rows
    def delete_excess_refs(starting_rows, required_rows, paragraphs):
        rows_to_remove = starting_rows - required_rows
        for index in range(rows_to_remove):
            delete_paragraph(paragraphs[-(index+1)])
        # number required => len(ref_order)
        # total at beginning => ref_table_paras
    delete_excess_refs(starting_rows=len(ref_table_paras), required_rows=len(ref_order), paragraphs=ref_table_paras)
    # use length of ref_list
    
    len(ref_order)
    def find_paragraghs_containing_string(target_string):
        return [para.text for para in document.paragraphs if target_string in para.text]
    # find [REF]
    # use index or filter where "[REF]" in para.text 
    # ref_paragraphs = find_paragraghs_containing_string(target_string="[REF]")

    # TODO: find reference table
    # extract references from csv
    # style appropriately
    # order from ref_order list
    # remove additional

    def alter_table_rows(total_rows, table, header_rows = 1):
        current_rows = len(table.rows)
        required_rows = total_rows + header_rows
        # remove bottom rows
        rows_to_remove = current_rows - required_rows
        # loop remove row[-1] * rows_to_remove
        for i in range(rows_to_remove):
            Delete_row_in_table(table, row=-1)

    def is_scenario_firefighting(name):
        if 'FSA'in name:
            return True

    def reformat_table_cell(cell):
        paragraphs = cell.paragraphs
        paragraphs[0].alignment = 1 # 1 = centered
        # for run in paragraphs.runs:
        run = paragraphs[0].runs
        font = run[0].font
        font.size = Pt(9) # pull from object
        font.name = font_name_light
        font.color.rgb = RGBColor(64,64,64) #gray color # to be pulled from object



    scenario_table = document.tables[2] # later find tables by name below??
    alter_table_rows(total_rows=len(scenario_names), table=scenario_table)
    # loop through scenarios

    # loop through each scenario in list
    def scenario_table_row(table_object, index_scenario):
        # column zero already has numbers
        model_name = scenario_names[index_scenario]
        row_index = index_scenario + 1
        row_object = table_object.rows[row_index]
        row_cells = row_object.cells
        num_cells = len(row_object.cells)

        venting_obj =scenarios_object[model_name]["venting"]
        mech_vent_obj = venting_obj["mech_extract"]
        for column_index in range(1, num_cells):
            # column 0 should be populated with numbers
        # column 1 - type
            if column_index == 1:
                is_firefighting = is_scenario_firefighting(model_name)
                if is_firefighting:
                    cell_text = 'Fire Service Access'
                if not is_firefighting:
                    cell_text = 'Means of Escape'
                replace_table_cell_content(cell=row_cells[column_index], replacement_text=cell_text) # how is it getting to other rows?
        # column 2 Reason for Modelling - for engineer
        # column 3 mech extract rate
            if column_index ==3:
                # model_name
                total_extract = (mech_vent_obj["number"] * mech_vent_obj["flow"])

                total_extract= round_to(value=total_extract)
                cell_text = str(total_extract)

                replace_table_cell_content(cell=row_cells[column_index], replacement_text=cell_text)

            # if column_index ==4:
            if column_index ==4:
                cell_text_list = [] # join with semi colons -> cell_text
                # TODO: how to superscript??
                total_supply = round_to(venting_obj["mech_supply"]["number"] * venting_obj["mech_supply"]["flow"])
                if total_supply:
                    cell_text_list.append(f'Mechanical Supply – {total_supply} m3/s')
                # check if above not zero
                # Mechanical Supply – 3.3 m3/s
                # 1.5m2 AOV; 1.5m2 Natural Smoke Shaft, 1.0m2 Corridor Vent; Though Stair Door via 1.0m2 AOV


                aov = venting_obj["stair_aov"]["area"]
                if aov:
                    cell_text_list.append(f'{aov} m2 AOV') # how to know if aov provided for head of stair or corridor?

                natural_inlet = sum(venting_obj["natural_openings"]) # check with Sam with artificial/for modelling inlet to be included?
                natural_inlet = round_to(natural_inlet)
                # if natural_inlet:
                #     cell_text_list.append(f'Additional inlet modelled in room of origin – {natural_inlet} m2') # how to superscript??

                cell_text = "; ".join(cell_text_list)

                replace_table_cell_content(cell=row_cells[column_index], replacement_text=f'ENGINEER TO CONFIRM: {cell_text}')
                # how to get type of supply??
        # column 4 inlet type

    for index_scenario in range(len(scenario_names)):
        scenario_table_row(table_object=scenario_table, index_scenario=index_scenario)

    def locate_table_from_cell(cell_row_index, cell_column_index, cell_text):
        table_list = []
        
        for table in document.tables:
            # check table has enough rows and columns
            cell = table.rows[cell_row_index].cells[cell_column_index]
            if cell_text in cell.text:
                table_list.append(table)
        return table_list
        # return all occurrences
    def return_general_scen(firefighting=False):
        if firefighting:
            return scenarios_object[FSA_scenarios[0]]
        else:
            return scenarios_object[MoE_scenarios[0]]
    def scenario_timeline_table(table_object, firefighting=False):
        table_rows = table_object.rows
        # loop through rows
        # if ff = False use first scenario for opening times
        # else use last
        scen_for_timings = return_general_scen(firefighting)

        for row_index in range(len(table_rows)):
            # first cell text
            row_cells = table_rows[row_index].cells
            row_title_cell = row_cells[0]
            if "Apartment Door Open" in row_title_cell.text:
                # column index 1 => stair door opening - add to opening object
                target_cell = table_rows[row_index].cells[1]
                cell_text = str(scen_for_timings["door_opening_times"]["opening_apartment"])
                replace_table_cell_content(cell=target_cell, replacement_text=cell_text)
            if "Stair Door Open" in row_title_cell.text:
                # column index 1 => stair door opening - add to opening object
                target_cell = table_rows[row_index].cells[1]
                # TODO: needs to be rounded
                cell_text = str(scen_for_timings["door_opening_times"]["opening_stair"])
                replace_table_cell_content(cell=target_cell, replacement_text=cell_text)
            if "Apartment Door Close" in row_title_cell.text:
                # column index 1 => stair door opening - add to opening object
                target_cell = table_rows[row_index].cells[1]
                cell_text = str(scen_for_timings["door_opening_times"]["closing_apartment"])
                replace_table_cell_content(cell=target_cell, replacement_text=cell_text)
                #  likely not needed column index 3 => scope whether stair door closes or not
                # probably leave max TD  and xxs based on walking speed of 1.2m/s
            if "Stair Door Close" in row_title_cell.text:
                target_cell = table_rows[row_index].cells[1]
                cell_text = str(scen_for_timings["door_opening_times"]["closing_stair"])
                replace_table_cell_content(cell=target_cell, replacement_text=cell_text)
                # if row_index == 7:
                # stair door closing -> add to opening object
                # model t
            if "Terminate" in row_title_cell.text:
                target_cell = table_rows[row_index].cells[1]
                cell_text = str(scen_for_timings["end_time"])
                replace_table_cell_content(cell=target_cell, replacement_text=cell_text)
                # model max t
        # first change @ row index 4

        # take timings from object
        # need maxTime in model
                
    MoE_FSA_tables = locate_table_from_cell(0,0,"Event")
    # Jinja removes timeline tables and sections when no FSA or no MOE present
    

    def populate_results_section(firefighting=False):
#     #     # bullet points where only 1 scen for fsa or MOE in section above
                #     #     # using jinja 

            results_tables = locate_table_from_cell(0,-1,"Meets Performance")
            if not firefighting: # i.e. MoE
                # if > 1 in scenarios
                moe_results_table = results_tables[0] 
                alter_table_rows(total_rows=len(MoE_scenarios), table=moe_results_table)
                table_rows = moe_results_table.rows
                for index in range(len(MoE_scenarios)):
                # use index for row index
                    row_index = index + 1 # top row for headings
                    row_cells = table_rows[row_index].cells
                    

                    scenario = MoE_scenarios[index]
                    # obtain from helper
                    tenable_time, max_pressure_drop, meet_criteria = scen_results_values(scenario, scenarios_object, firefighting)
                    # for bullets set varible to above values
                    replace_table_cell_content(cell=row_cells[1], replacement_text=str(round_to(tenable_time)))

                    replace_table_cell_content(cell=row_cells[2], replacement_text=f'{max_pressure_drop}kPa')
                    replace_table_cell_content(cell=row_cells[3], replacement_text=meet_criteria)
            elif firefighting:
                fsa_results_table = results_tables[-1] # if > 1 in scenarios; else bullet
                alter_table_rows(total_rows=len(FSA_scenarios), table=fsa_results_table, header_rows=2)
                table_rows = fsa_results_table.rows                   
                #     if len(fsa_or_moe_scenarios) == 1:
                for index in range(len(FSA_scenarios)):
                    # use index for row index
                        row_index = index + 2 # top 2 row for headings; double row
                        row_cells = table_rows[row_index].cells
                        

                        scenario = FSA_scenarios[index]
                        tenable_object = scenarios_object[scenario]["tenability"] # applies to bullets
                        tenability_keys = list(tenable_object.keys())   # this and below applies to bullets
                        text_list, worst_temp, worst_vis, meet_criteria, max_pressure_drop = scen_results_values(scenario, scenarios_object, firefighting)
                        for index_key in range(len(text_list)):
                            replace_table_cell_content(cell=row_cells[index_key+1], replacement_text=(text_list[index_key]))

                        replace_table_cell_content(cell=row_cells[4], replacement_text=str(round_to(worst_vis)))
                        replace_table_cell_content(cell=row_cells[5], replacement_text=str(round_to(worst_temp)))
                        replace_table_cell_content(cell=row_cells[-2], replacement_text=str((max_pressure_drop)))
                        replace_table_cell_content(cell=row_cells[-1], replacement_text=meet_criteria)

        
    # if MoE >0; then will be first table
    if len(MoE_scenarios) > 0:
        moe_table = MoE_FSA_tables[0]
        scenario_timeline_table(table_object=moe_table, firefighting=False)            
        if len(MoE_scenarios) > 1: # if moe > 1; else would be bullet points      
            populate_results_section(firefighting=False)
    if len(FSA_scenarios) > 0:
        fsa_table = MoE_FSA_tables[-1] 
        scenario_timeline_table(table_object=fsa_table, firefighting=True)
        if len(FSA_scenarios) > 1: # if > 1 use table for results, else populate bullets
            populate_results_section(firefighting=True)

    document.save(output_path)

